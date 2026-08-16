
"""   Generating user story for given requirement and updating KB and generating test cases for requested user story """

import os
import re
from typing import List, Tuple

from langchain_community.document_loaders import Docx2txtLoader
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_core.prompts import PromptTemplate
from docx import Document  # pip install python-docx

class RAGModel:
    def __init__(self, kb_path: str, google_api_key: str):
        self.kb_path = kb_path
        self.google_api_key = google_api_key

        # Configure embeddings & LLM
        self.embeddings = GoogleGenerativeAIEmbeddings(
            model="gemini-embedding-2",
            google_api_key=self.google_api_key,
        )
        self.llm = ChatGoogleGenerativeAI(
            model="gemini-flash-latest",
            temperature=0.3,
            google_api_key=self.google_api_key,
        )
        self.vector_db = None
        self.retriever = None

        # Build FAISS from KB
        self._build_vector_db()

    def _build_vector_db(self):
        if not os.path.exists(self.kb_path):
            raise FileNotFoundError(f"KB file not found: {self.kb_path}")

        docs = Docx2txtLoader(self.kb_path).load()       # Load the Word document and extract text content
        splitter = RecursiveCharacterTextSplitter(      # Create text splitter
            chunk_size=1800,
            chunk_overlap=300,
            separators=["\n## ", "\n### ", "\n", ". ", " ", ""],   # Uses hierarchical separators (headers → paragraphs → sentences → words)

        )

        texts = []
        for d in docs:
            texts.extend(splitter.split_text(d.page_content))   # Split each document into smaller chunks and collect all text pieces

        self.vector_db = FAISS.from_texts(texts, embedding=self.embeddings)     #Create FAISS vector database from text chunks using embeddings
        self.vector_db.save_local("index_db")      # Save database locally for persistence
        self.retriever = self.vector_db.as_retriever(search_kwargs={"k": 6})    # Create retriever that returns top 6 most similar chunks

    def rebuild_index(self):
        """Rebuild FAISS after KB update."""
        self._build_vector_db()

    def _retrieve(self, query: str) -> List[str]:
        docs = self.retriever.invoke(query)     # Private method to retrieve relevant documents for a query
        return [d.page_content for d in docs]   # Returns list of text content from retrieved documents

    @staticmethod
    def _build_prompt(user_req: str, kb_context: List[str], include_neg: bool, refinements: str) -> PromptTemplate:
        context_text = "\n\n".join(kb_context)  # Join retrieved context with double newlines
        neg = "- Include negative/edge scenarios and failure handling.\n" if include_neg else "- Focus on happy-path; negative scenarios optional.\n"      # Set negative scenario instruction based on flag
        ref = f"\nAdditional refinements:\n{refinements}\n" if refinements.strip() else ""          # Add refinements section if provided
        template = f"""
You are an AI agent that generates **structured Agile user stories** with **acceptance criteria**.
Leverage the retrieved **knowledge base** (templates, compliance rules, historical examples, numbered requirements).
Respond in **Markdown**.

### Requirement (from user)
{{user_req}}

### Retrieved KB Context (RAG)
{context_text}

### Instructions
- Use format: *As a <role>, I want <capability>, so that <benefit>*.
- 2–3 stories per feature; each ≤ 120 words.
- **Acceptance Criteria** in **Gherkin** (Given/When/Then), 3–5 bullets per story.
- Integrate applicable **Business Rules** and **Compliance** (security baseline, GDPR/CCPA, WCAG 2.1 AA, PCI DSS, audit/retention, etc.).
{neg}- Keep output concise, precise, implementation-agnostic.
{ref}

### Output
1) **User Stories**
2) **Acceptance Criteria** (per story)
3) **Non-Functional Criteria**
4) **Standards Validation Summary** (gaps)
"""
        return PromptTemplate(template=template, input_variables=["user_req"])

    def generate(
        self,
        user_req: str,
        include_neg: bool,
        refinements: str = "",
        extra_query_terms: List[str] = None,
    ) -> str:
        augmented_query = user_req
        if extra_query_terms:
            augmented_query = f"{user_req}\n" + " ".join(extra_query_terms)

        kb_context = self._retrieve(augmented_query)
        prompt = self._build_prompt(user_req, kb_context, include_neg, refinements)
        chain = prompt | self.llm
        result = chain.invoke({"user_req": user_req})

        processed_content = result.content if isinstance(result.content, str) else " ".join(map(str, result.content))
        return processed_content

    # ---------- Corrective RAG (unchanged minimal) ----------
    def _validate_output(self, output: str, include_neg: bool) -> Tuple[bool, List[str]]:
        gaps = []
        required_sections = [
            "User Stories", "Acceptance Criteria", "Non-Functional Criteria", "Standards Validation Summary",
        ]
        for sec in required_sections:
            if sec.lower() not in output.lower():
                gaps.append(f"Missing section: **{sec}**")

        story_lines = re.findall(r"(?:^|\n)\s*(?:\*?\s*)As a .*?, I want .*?, so that .*?(\.|$)", output, flags=re.IGNORECASE)
        if not (2 <= len(story_lines) <= 3):
            gaps.append("Ensure **2–3 user stories** using the format: *As a <role>, I want <capability>, so that <benefit>*.")

        story_line_texts = re.findall(r"(?:^|\n)\s*(?:\*?\s*)(As a .*?, I want .*?, so that .*?)(?:\.|$)", output, flags=re.IGNORECASE)
        for s in story_line_texts:
            words = re.findall(r"\w+", s)
            if len(words) > 120:
                gaps.append("Each **user story** should be **≤ 120 words**. Shorten long stories.")

        if not re.search(r"\bGiven\b", output, re.IGNORECASE) or not re.search(r"\bWhen\b", output, re.IGNORECASE) or not re.search(r"\bThen\b", output, re.IGNORECASE):
            gaps.append("Add **Gherkin-style Acceptance Criteria** with **Given/When/Then** (3–5 bullets per story).")

        compliance_terms = ["GDPR", "CCPA", "WCAG", "PCI", "security", "audit", "retention"]
        if not any(term.lower() in output.lower() for term in compliance_terms):
            gaps.append("Explicitly reference **GDPR**, **CCPA**, **WCAG 2.1 AA**, **PCI DSS**, **security baseline**, **audit & retention** in criteria or summary.")

        if include_neg:
            neg_keywords = ["negative", "edge", "error", "failure", "unauthorized", "invalid", "timeout"]
            if not any(k in output.lower() for k in neg_keywords):
                gaps.append("Include **negative/edge scenarios** and **failure handling** in acceptance criteria.")

        return len(gaps) == 0, gaps

    def _build_refinements_from_gaps(self, gaps: List[str]) -> Tuple[str, List[str]]:
        refinements = "Please correct the following:\n" + "\n".join([f"- {g}" for g in gaps])
        extra_terms = []
        glist = [g.lower() for g in gaps]
        if any("gdpr" in g or "ccpa" in g for g in glist): extra_terms += ["GDPR", "CCPA", "privacy", "data protection", "retention"]
        if any("wcag" in g for g in glist): extra_terms += ["WCAG 2.1 AA", "accessibility", "screen reader", "contrast"]
        if any("pci" in g for g in glist): extra_terms += ["PCI DSS", "payment", "cardholder", "encryption"]
        if any("security" in g for g in glist): extra_terms += ["security baseline", "authentication", "authorization", "audit"]
        if any("gherkin" in g or "given/when/then" in g for g in glist): extra_terms += ["Given When Then", "acceptance criteria", "Gherkin"]
        if any("negative" in g or "edge" in g or "failure" in g for g in glist): extra_terms += ["error handling", "edge cases", "invalid input", "unauthorized", "timeout"]
        # Dedup
        seen, uniq = set(), []
        for t in extra_terms:
            if t not in seen:
                uniq.append(t); seen.add(t)
        return refinements, uniq

    def corrective_generate(self, user_req: str, include_neg: bool, initial_refinements: str = "", max_loops: int = 3) -> str:
        refinements, extra_terms, last_output = initial_refinements or "", [], ""
        for i in range(1, max_loops + 1):
            output = self.generate(user_req, include_neg, refinements, extra_terms)
            last_output = output
            passed, gaps = self._validate_output(output, include_neg)
            print(f"\n[Validation Iteration {i}] Passed: {passed}")
            for g in gaps: print(f" - {g}")
            if passed:
                print("\n[Corrective RAG] Validation passed.")
                return output
            refine_msg, new_terms = self._build_refinements_from_gaps(gaps)
            refinements = (refinements + "\n" + refine_msg).strip() if refinements else refine_msg
            for t in new_terms:
                if t not in extra_terms:
                    extra_terms.append(t)
            print("\n[Corrective RAG] Regenerating with refinements & augmented retrieval...")
        print("\n[Corrective RAG] Max loops reached. Returning last output with potential gaps.")
        return last_output

    # ---------- NEW: KB update + test case generation ----------
    @staticmethod
    def _extract_user_stories(output: str) -> List[str]:
        """
        Extract user stories from the 'User Stories' section.
        """
        # Get the section between 'User Stories' and 'Acceptance Criteria'
        m = re.search(r"(?si)User Stories.*?\n(.*?)(?:\n\s*\d\)\s*\*\*Acceptance Criteria\*\*|\n\s*Acceptance Criteria|\Z)", output)
        block = m.group(1).strip() if m else output

        # Collect lines starting with 'As a ...'
        stories = re.findall(r"(?:^|\n)\s*(?:[-*]\s*)?(As a .*?, I want .*?, so that .*?)(?:\.|\n|$)", block, flags=re.IGNORECASE)
        # Normalize spacing and add trailing dot
        stories = [s.strip().rstrip(".") + "." for s in stories]
        return stories

    def append_user_stories_to_kb(self, user_req: str, stories: List[str]) -> None:
        """
        Append the generated user stories to the KB .docx with a simple heading.
        """
        if os.path.exists(self.kb_path):
            doc = Document(self.kb_path)
        else:
            doc = Document()
        doc.add_heading("Generated User Stories", level=1)
        doc.add_paragraph(f"Requirement: {user_req}")
        for i, s in enumerate(stories, start=1):
            doc.add_paragraph(f"{i}. {s}")
        doc.save(self.kb_path)
        print(f"[KB Update] Appended {len(stories)} user stories to: {self.kb_path}")

    @staticmethod
    def _build_test_prompt(story_text: str, kb_context: List[str], include_neg: bool) -> PromptTemplate:
        ctx = "\n\n".join(kb_context)
        neg_line = "- Include negative/edge scenarios and failure handling.\n" if include_neg else "- Negative scenarios optional.\n"
        template = f"""
You are a QA test designer. Use the KB context and the selected **user story** to generate concise, high-quality **test cases**.

### Selected User Story
{{story_text}}

### Retrieved KB Context (RAG)
{ctx}

### Instructions
- Provide a numbered list of **test cases** with: Title, Preconditions, Test Data (if any), Steps, Expected Result.
- Include **Gherkin scenarios** (Given/When/Then), 4–6 total, mapped to the story.
- Align with relevant **Business Rules, Compliance**, and **Non-Functional** criteria referenced in context.
{neg_line}- Keep output concise and implementation-agnostic.

### Output
1) **Test Cases** (numbered)
2) **Gherkin Scenarios**
"""
        return PromptTemplate(template=template, input_variables=["story_text"])

    def generate_test_cases(self, story_text: str, include_neg: bool) -> str:
        """
        Fetch context from updated KB and generate test cases for the given user story.
        """
        kb_context = self._retrieve(story_text + " test cases acceptance criteria compliance performance accessibility")
        prompt = self._build_test_prompt(story_text, kb_context, include_neg)
        chain = prompt | self.llm
        res = chain.invoke({"story_text": story_text})
        return res.content if isinstance(res.content, str) else " ".join(map(str, res.content))


def main():
    print("=== AI-Powered User Story & Acceptance Criteria Generator (RAG + Corrective + KB Update + Test Cases) ===")

    
     # Get Google API key from environment if already set; otherwise prompt user
    google_key = os.environ.get("GOOGLE_API_KEY")
    if not google_key:
        google_key = input("Enter your Google API Key: ").strip()
    if not google_key:
        raise ValueError("Google API key is required.")
    os.environ["GOOGLE_API_KEY"] = google_key
    
    # Prompt for KB path dynamically in VS Code
    kb_path = input("Enter KB .docx path (example: C:\\Users\\paila\\OneDrive\\Desktop\\digital_wallet_business_compliance_rules.docx): ").strip()
    
    # Remove accidental quotes if user pastes a path with quotes
    kb_path = kb_path.strip('"').strip("'")
    
    if not kb_path:
        print("No file path entered. Please provide a valid .docx path.")
        raise ValueError("KB file path is required.")

    rag = RAGModel(kb_path=kb_path, google_api_key=google_key)

    user_req = input("Enter the requirement (e.g., 'Unified Login', 'RBAC', or a freeform description): ").strip()
    include_neg = input("Include negative scenarios? (y/n): ").strip().lower().startswith("y")

    use_corrective = input("Enable corrective RAG (auto revalidate & regenerate)? (y/n): ").strip().lower().startswith("y")
    max_loops_str = input("Max corrective loops (default 3): ").strip()
    max_loops = int(max_loops_str) if max_loops_str.isdigit() else 3

    # Generate
    if use_corrective:
        output = rag.corrective_generate(user_req=user_req, include_neg=include_neg, initial_refinements="", max_loops=max_loops)
    else:
        output = rag.generate(user_req=user_req, include_neg=include_neg, refinements="")

    print("\n--- Generated Output ---\n")
    print(output)

    # Extract user stories and update KB
    stories = rag._extract_user_stories(output)
    if not stories:
        print("\n[KB Update] No user stories detected; skipping KB update.")
    else:
        rag.append_user_stories_to_kb(user_req, stories)
        # Rebuild FAISS from updated KB for downstream test case generation
        rag.rebuild_index()

    # Ask for test case generation
    if stories:
        want_tests = input("\nDo you want to generate test cases for any generated user story? (y/n): ").strip().lower().startswith("y")
        if want_tests:
            print("\nSelect story indices (comma-separated).")
            for i, s in enumerate(stories, start=1):
                print(f"{i}) {s}")
            sel = input("Enter indices (e.g., 1,2): ").strip()
            # Parse indices
            try:
                idxs = [int(x.strip()) for x in sel.split(",") if x.strip().isdigit()]
            except Exception:
                idxs = []

            if not idxs:
                print("No valid indices provided. Exiting test case generation.")
            else:
                print("\n=== Test Cases ===")
                for i in idxs:
                    if 1 <= i <= len(stories):
                        tc = rag.generate_test_cases(stories[i-1], include_neg=include_neg)
                        print(f"\n--- Story {i} Test Cases ---\n")
                        print(tc)
                    else:
                        print(f"\nIndex {i} out of range; skipped.")
        else:
            print("Skipping test case generation.")

    print("\nDone.")

if __name__ == "__main__":
    main()
